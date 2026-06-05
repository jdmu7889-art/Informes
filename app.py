import streamlit as st
from docx import Document
from docx.shared import Inches
import io

# Configuración de la página
st.set_page_config(page_title="Generador de Informes de Mantenimiento", layout="wide")

st.title("📊 Generador de Informes de Mantenimiento Fotovoltaico")
st.write("Llena los datos del campo y captura las fotografías solicitadas. Al finalizar, descarga tu informe en Word.")

# --- SECCIÓN 1: DATOS DEL INFORME ---
st.header("📝 Datos Generales")
col1, col2, col3 = st.columns(3)
with col1:
    estacion = st.text_input("Nombre de la Estación")
with col2:
    contratista = st.text_input("Contratista")
with col3:
    fecha = st.date_input("Fecha de Mantenimiento")

# --- SECCIÓN 2: CONFIGURACIÓN DE FOTOGRAFÍAS REQUERIDAS ---
# Definimos la lista de requerimientos según tu listado
requerimientos_fotos = [
    {"id": "permiso_alturas", "label": "Fotografía de Permiso de Trabajo en Alturas", "cant": 1},
    {"id": "permiso_general", "label": "Fotografía del Permiso General de Trabajo", "cant": 1},
    {"id": "paneles_previo", "label": "Paneles Previo a Limpieza", "cant": 2},
    {"id": "paneles_limpios", "label": "Paneles Limpios", "cant": 2},
    {"id": "inversores_previo", "label": "Equipos Inversores y Paneles Eléctricos Previo a Limpieza", "cant": 4},
    {"id": "inversores_limpios", "label": "Equipos Inversores y Paneles Eléctricos Limpios", "cant": 4},
]

# Agregar dinámicamente las mediciones DC (S1 a S10)
for i in range(1, 11):
    requerimientos_fotos.append({"id": f"corr_dc_s{i}", "label": f"Fotografía Medición de Corriente DC S{i}", "cant": 1})
for i in range(1, 11):
    requerimientos_fotos.append({"id": f"volt_dc_s{i}", "label": f"Fotografía Medición de Voltaje DC S{i}", "cant": 1})

# Agregar mediciones AC Fases y Breakers
for i in range(1, 4):
    requerimientos_fotos.append({"id": f"corr_ac_f{i}", "label": f"Fotografía Medición de Corriente AC Fase {i}", "cant": 1})
for i in range(1, 4):
    requerimientos_fotos.append({"id": f"corr_ac_b{i}", "label": f"Fotografía Medición de Corriente AC Breaker {i}", "cant": 1})

# Resocados y Termografías
requerimientos_fotos.extend([
    {"id": "resocado_ac", "label": "Fotografía De Resocado Panel AC", "cant": 1},
    {"id": "resocado_dc", "label": "Fotografía De Resocado Panel DC", "cant": 1},
    {"id": "resocado_solares", "label": "Fotografía De Resocado PANELES SOLARES", "cant": 1},
    {"id": "termo_paneles", "label": "Fotografía Termográfica de Paneles", "cant": 1},
    {"id": "termo_b_ac", "label": "Fotografía Termográfica de Breaker AC", "cant": 1},
    {"id": "termo_b_dc", "label": "Fotografía Termográfica de Breaker DC", "cant": 1},
])

# --- SECCIÓN 3: CAPTURA DE FOTOS EN LA UI ---
st.header("📸 Registro Fotográfico Obligatorio")
st.info("Puedes usar la cámara de tu dispositivo o subir archivos guardados.")

fotos_guardadas = {}

# Iteramos por cada requerimiento para crear los botones de carga
for req in requerimientos_fotos:
    st.subheader(req["label"])
    fotos_guardadas[req["id"]] = []
    
    # Si requiere más de una foto, creamos columnas
    cols = st.columns(req["cant"])
    for i in range(req["cant"]):
        with cols[i]:
            label_foto = f"Captura {i+1} de {req['cant']}" if req["cant"] > 1 else "Tomar Foto"
            foto = st.camera_input(label_foto, key=f"{req['id']}_{i}")
            # Si no hay cámara activa, permitimos subir archivo alternativo
            if foto is None:
                foto = st.file_uploader(f"O subir imagen {i+1}", type=["jpg", "jpeg", "png"], key=f"file_{req['id']}_{i}")
            
            if foto is not None:
                fotos_guardadas[req["id"]].append(foto)

# --- SECCIÓN 4: GENERACIÓN DEL DOCUMENTO WORD ---
st.markdown("---")
if st.button("🚀 Generar y Descargar Informe Word", type="primary"):
    if not estacion or not contratista:
        st.error("⚠️ Por favor ingresa el Nombre de la Estación y el Contratista antes de generar el informe.")
    else:
        # Crear documento Word en memoria
        doc = Document()
        
        # Título del Informe
        doc.add_heading("INFORME TÉCNICO DE MANTENIMIENTO FOTOVOLTAICO", level=0)
        
        # Tabla de Datos Generales
        doc.add_heading("1. Datos Generales", level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Shading Accent 1'
        
        table.cell(0, 0).text = "Estación:"
        table.cell(0, 1).text = estacion
        table.cell(1, 0).text = "Contratista:"
        table.cell(1, 1).text = contratista
        table.cell(2, 0).text = "Fecha de Mantenimiento:"
        table.cell(2, 1).text = str(fecha)
        
        doc.add_paragraph("\n")
        doc.add_heading("2. Registro Fotográfico y Evidencias", level=1)
        
        # Insertar Fotos de manera ordenada
        for req in requerimientos_fotos:
            doc.add_heading(req["label"], level=2)
            lista_fotos = fotos_guardadas[req["id"]]
            
            if lista_fotos:
                for idx, foto_data in enumerate(lista_fotos):
                    # Convertir los datos cargados en bytes para python-docx
                    foto_bytes = io.BytesIO(foto_data.getvalue())
                    doc.add_paragraph(f"Evidencia {idx + 1}:")
                    doc.add_picture(foto_bytes, width=Inches(4.5)) # Ajusta el tamaño de la imagen en el Word
            else:
                doc.add_paragraph("⚠️ No se adjuntó evidencia fotográfica para este punto.")
                
        # Guardar en buffer de memoria para la descarga
        target_stream = io.BytesIO()
        doc.save(target_stream)
        target_stream.seek(0)
        
        st.success("¡Informe generado con éxito!")
        st.download_button(
            label="💾 Descargar archivo Word",
            data=target_stream,
            file_name=f"Informe_Mantenimiento_{estacion.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )